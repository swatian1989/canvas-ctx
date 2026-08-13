"""Assembles reports/analysis_report.{md,html,docx} from figures.py + tables.py.

`build_report()` is the single entry point scripts/run_report.py calls. It
runs every figure/table function (cheap: they read cached artefacts, they do
not re-run training or feature extraction -- see the module docstrings in
figures.py/tables.py), then renders three parallel views of the SAME content
so nothing drifts between formats:

    reports/analysis_report.md    plain markdown, images as relative links
    reports/analysis_report.html  self-contained, images inlined as base64
    reports/analysis_report.docx  navy/steel blue, Calibri, justified body

The report's content -- every paragraph of prose -- lives in `build_sections`
below as a list of `Section` objects. All three renderers just walk that
list, so editing report text means editing one place.
"""

from __future__ import annotations

import base64
import hashlib
import importlib.metadata as im
import json
import subprocess
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd

from . import figures as figs
from . import tables as tbls
from .style import NAVY, STEEL_BLUE

# Every third-party library the pipeline actually uses, with the role it plays.
# Listed exhaustively so the software environment is reproducible from the
# report alone, without reading the source.
SOFTWARE: list[tuple[str, str, str]] = [
    # (distribution, import name, role)
    ("numpy", "numpy", "Array computation throughout; all coordinate and "
                       "embedding maths"),
    ("pandas", "pandas", "Tabular data: single-cell tables, patch labels, "
                         "feature matrices"),
    ("scipy", "scipy", "cKDTree for 40 um neighbourhood and nearest-neighbour "
                       "queries; sparse matrices; Wilcoxon signed-rank test; "
                       "FFT cross-correlation for registration QC"),
    ("scikit-learn", "sklearn", "k-means, latent Dirichlet allocation, "
                                "silhouette / Davies-Bouldin / adjusted Rand "
                                "index, agglomerative consensus clustering, "
                                "precision/recall/F1 and confusion matrices"),
    ("pyyaml", "yaml", "Configuration files and the inherit-chain resolver"),
    ("pyarrow", "pyarrow", "Parquet engine for every cached artefact"),
    ("tqdm", "tqdm", "Progress reporting in long stage loops"),
    ("tabulate", "tabulate", "Markdown table rendering in the report"),

    ("lifelines", "lifelines", "Cox proportional hazards models, Kaplan-Meier "
                               "estimation, log-rank tests, concordance index"),
    ("scikit-survival", "sksurv", "LASSO-Cox (CoxnetSurvivalAnalysis), random "
                                  "survival forest permutation importance, "
                                  "time-dependent cumulative AUC"),
    ("statsmodels", "statsmodels", "Benjamini-Hochberg false discovery rate "
                                   "correction"),

    ("networkx", "networkx", "Connected-component fallback for collinearity "
                             "reduction"),
    ("python-igraph", "igraph", "Louvain community detection on the feature "
                                "correlation graph"),

    ("torch", "torch", "All neural network components: the CANVAS head, focal "
                       "loss, weighted sampling, and the graph/2D-grid/3D-grid "
                       "context encoders"),
    ("torchvision", "torchvision", "ResNet-50 benchmark encoder and its own "
                                   "preprocessing transforms"),
    ("timm", "timm", "Model registry and data configuration for UNI and MUSK"),
    ("transformers", "transformers", "Phikon encoder (ViTModel) and its "
                                     "AutoImageProcessor"),
    ("huggingface-hub", "huggingface_hub", "Weight download and token-gated "
                                           "model access"),

    ("pillow", "PIL", "Patch image handling and resampling"),
    ("opencv-python-headless", "cv2", "Tissue segmentation (HSV, Otsu, "
                                      "morphology), artefact filtering, "
                                      "marker gating thresholds"),
    ("tifffile", "tifffile", "Pyramidal OME-TIFF reading for the paired H&E "
                             "whole-slide images"),
    ("imagecodecs", "imagecodecs", "JPEG and zlib codecs those OME-TIFFs use"),
    ("zarr", "zarr", "Windowed level-0 reads, so a 13.5 GB image is never "
                     "loaded whole"),

    ("matplotlib", "matplotlib", "Every figure in this report"),
    ("python-docx", "docx", "Word rendering of the report and manuscript"),
    ("openpyxl", "openpyxl", "Reads the TCGA-CDR clinical workbook"),

    ("pytest", "pytest", "Test suite"),
]

PACKAGES = [dist for dist, _imp, _role in SOFTWARE]

# Deliberately optional. Each has a documented fallback so the pipeline runs
# without it, and the fallback taken is recorded.
OPTIONAL_SOFTWARE: list[tuple[str, str]] = [
    ("spatial-lda", "True spatial-LDA prior with a smoothness penalty over "
                    "neighbouring index cells. NOT installed in this run, so "
                    "scikit-learn's LatentDirichletAllocation was used instead: "
                    "the same generative model without the spatial prior. This "
                    "is a real deviation and is recorded in T10."),
    ("openslide-python", "Aperio/SVS reading for TCGA slides. Not required for "
                         "the Orion cohort, whose OME-TIFFs are read with "
                         "tifffile and zarr."),
    ("stardist", "Nuclear segmentation. Not required here: the paired cohort "
                 "ships its own segmented single-cell tables."),
]


@dataclass
class Section:
    heading: str
    level: int
    paragraphs: list[str] = field(default_factory=list)
    figure_ids: list[str] = field(default_factory=list)
    table_ids: list[str] = field(default_factory=list)


# ============================================================== run everything


def _run_all(figures_dir: str, tables_dir: str) -> tuple[dict, dict]:
    fig = {m["id"]: m for m in (fn(figures_dir) for fn in figs.ALL_FIGURES)}
    tab = {m["id"]: m for m in (fn(tables_dir) for fn in tbls.ALL_TABLES)}
    return fig, tab


def _reproducibility_stats(config_path: str = "config/crc_train_brca_apply.yaml") -> dict:
    import sys
    sys.path.insert(0, "src")
    from canvas_brca.utils.config import load_config

    cfg = load_config(config_path)
    cfg_hash = hashlib.sha256(json.dumps(cfg, sort_keys=True, default=str).encode()).hexdigest()

    versions = {}
    for pkg in PACKAGES:
        try:
            versions[pkg] = im.version(pkg)
        except im.PackageNotFoundError:
            versions[pkg] = "not installed"

    try:
        sha = subprocess.run(["git", "rev-parse", "HEAD"], capture_output=True, text=True,
                            timeout=5).stdout.strip()
        git_sha = sha if sha else "NOT A GIT REPOSITORY"
    except Exception:
        git_sha = "NOT A GIT REPOSITORY"

    import platform
    return {"config_path": config_path, "config_hash_sha256": cfg_hash[:16],
           "versions": versions, "git_sha": git_sha,
           "python_version": platform.python_version(),
           "platform": f"{platform.system()} {platform.release()}",
           "project_seed": cfg["project"]["seed"]}


# ==================================================================== content


def build_sections(fig: dict, tab: dict, stats: dict) -> list[Section]:
    n_placeholder_fig = sum(1 for f in fig.values() if f["source"] == "MISSING DATA")
    n_placeholder_tab = sum(1 for t in tab.values() if t["source"] == "MISSING DATA")

    t5 = tab["T5"]["df"]
    base_mode = t5["mode"].iloc[0]
    grid3d_row = t5[(t5["mode"] == "grid3d") & (t5["metric"] == "macro_f1")]
    none_row = t5[(t5["mode"] == "none") & (t5["metric"] == "macro_f1")]
    n_seeds_main = int(t5["n_seeds"].iloc[0])

    sections = [
        Section("Summary", 1, [
            "This report covers work that needed no downloaded data: `data/raw/` is "
            "confirmed empty (checked at the start of this run), so every quantitative "
            "result below comes from one of two sources, and every figure/table caption "
            "says which. First, a synthetic `--simulate` fixture built into "
            "`scripts/run_stage5_features.py` and `scripts/run_final_benchmark.py`, used "
            "throughout the codebase to test the statistical and modelling machinery "
            "before real predictions exist. Second, a null-outcome calibration check: the "
            "same simulated 262-feature matrix joined to an independently random survival "
            "outcome, to confirm the stage 6 statistics do not hallucinate significance.",
            f"Also completed this run: the previously-missing embedding-extraction pipeline "
            f"(`stage3_model/encoders.py`, `stage3_model/extract_embeddings.py`, "
            f"`scripts/run_stage3_extract.py`) and two new runner scripts "
            f"(`scripts/run_stage2_pair.py`, `scripts/run_stage4_infer.py`), plus this "
            f"reporting package, all with tests against synthetic fixtures and zero network "
            f"calls. The test suite is at **66 passing** (39 baseline + 20 embedding/encoder "
            f"+ 7 reporting), never dropped below the 39 floor.",
            f"Headline numbers, all SIMULATED (see caveats): the 6-seed, 4-mode context "
            f"benchmark shows mode=none (Method 1/CANVAS) at "
            f"{none_row['mean'].iloc[0]:.3f} macro-F1 vs grid3d at "
            f"{grid3d_row['mean'].iloc[0]:.3f}, every context mode significant against "
            f"mode={base_mode} at the {n_seeds_main}-seed paired-Wilcoxon floor (p=0.0312). "
            f"The null-calibration Cox check is clean (0/261 features FDR-significant, "
            f"n=200) but the downstream multivariable signature model is NOT clean "
            f"(C-index 0.60 at n=200, worse at n=50) -- see Section 3.6 and the "
            f"Limitations.",
            "**Three caveats that matter most:** (1) zero real tissue or clinical data "
            f"was used anywhere in this report -- {n_placeholder_fig} of 22 figures and "
            f"{n_placeholder_tab} of 10 tables are labelled MISSING DATA placeholders, not "
            "filled with anything approximate. (2) The benchmark's synthetic fixture has "
            "oriented bands built in specifically to reward spatial context (documented in "
            "the script itself); a large simulated gain there is a machinery check, never "
            "evidence about real tissue. (3) The stage 6 signature model shows real, "
            "reproducible in-sample optimism bias -- confirmed across independent seeds and "
            "sample sizes -- that will inflate any real hazard ratio or C-index reported "
            "from Phase 5/6 unless corrected with a held-out split before publication.",
            "F1 shows the overall design and where it departs from the published "
            "protocol; T1 is the full data inventory, including what is still missing "
            "and what each missing cohort would unblock.",
        ], figure_ids=["F1"], table_ids=["T1"]),

        Section("Methods", 1, [
            "Every parameter below is taken from the CANVAS STAR Methods and marked "
            "[PAPER] in `PROTOCOL.md`; none was changed in this run. Where this project's "
            "design deviates from the published protocol -- substituted cohorts, a "
            "different encoder, reduced permutation counts, two DELIBERATE statistical "
            "corrections -- the deviation is flagged inline here and fully itemised in "
            "**T10** (parameter, paper value, our value, reason, expected impact).",
        ]),
        Section("CN discovery [PAPER: 40 um radius, k=5-20 sweep, silhouette/DB/adjacent-ARI]",
               2, ["40 um neighbourhood radius (~25 neighbours), spatial-LDA over local "
                  "composition, k-means sweep k=5-20 scored on silhouette, Davies-Bouldin "
                  "and adjacent-k ARI, final k=10. Cohort substitution: see T10 row 1."]),
        Section("Label transfer [PAPER: 5 um registration threshold, patch purity rules]",
               2, ["Affine registration, centroid-to-centroid threshold 5 um, patches "
                  "labelled background at <=5 cells, dominant CN at >15 cells and >=60% "
                  "composition, discarded otherwise. Split at the sample (patient) level, "
                  "never patch level. Cohort substitution: see T10 rows 1-2."]),
        Section("Model [PAPER: 224x224 patches, 256->128->K+1 head, focal loss + weighted sampling]",
               2, ["Head architecture, loss, and sampler are unchanged from PROTOCOL.md's "
                  "spec and are NOT touched by this project (`stage3_model/head.py` is "
                  "pre-existing, tested, untouched). Encoder and fine-tuning policy "
                  "deviations: T10 row 3."]),
        Section("Feature engineering [PAPER: 262 features, k=6 transition entropy]",
               2, ["Composition (10) + diversity (6) + dispersion (90) + interaction (100) "
                  "+ distance (55) + transition (1) = 262, asserted at runtime in "
                  "`spatial_features.py` (untouched, protected). Two DELIBERATE deviations "
                  "live here: toroidal-shift interaction null instead of label shuffling, "
                  "and border/Donnelly edge correction on dispersion metrics. Both are "
                  "reproduced and quantified in F21/F22 and itemised in T10 rows 6-7."]),
        Section("Clinical modelling [PAPER: univariate Cox+FDR, PAM/Canberra consensus clustering, LASSO-Cox to RSF to multivariable signature]",
               2, ["`stage6_clinical/signature.py` (untouched, protected) implements this "
                  "chain exactly as specified. This run exercises the full chain against "
                  "the null-outcome fixture (Section 3.6) -- the paper's method is not "
                  "changed, but its BEHAVIOUR under a null is characterised here, which the "
                  "paper's methods section does not itself require."]),

        Section("Results", 1, []),

        Section("3.1 CN discovery (Stage 1) -- REAL DATA", 2, [
            "The Schurch 2020 colorectal CODEX single-cell table (Mendeley mpjzbtfgfr, "
            "223 MB, sha256 verified, PMID 32763154) was downloaded and analysed. This is "
            "the only stage in this report built on real tissue measurements.",
            "**Coordinate units were established before any analysis, not assumed.** Raw "
            "X spans 0-1919 and Y spans 0-1439 in every one of the 140 images: a 1920x1440 "
            "sensor grid, so pixels. The Cancer Imaging Archive collection page for this "
            "exact dataset documents 377.44 nm/px on a Keyence BZ-X710 with a 20x/0.75 "
            "objective, giving a 724 x 543 um field of view. An independent check confirms "
            "it: median nearest-neighbour cell spacing comes out at 6.8 um, consistent with "
            "packed tissue, whereas the alternative 0.325 um/px figure quoted for other "
            "CODEX installations would compress that to 5.9 um and make the cells "
            "implausibly close.",
            "**The 40 um neighbourhood contains a median of 33 cells here, not the ~25 the "
            "CANVAS STAR Methods states.** This was diagnosed rather than absorbed. The "
            "~25 figure derives from the published NSCLC CODEX cohort, a different tissue; "
            "the colorectal invasive front is denser, at a median 4,890 cells/mm2. "
            "Critically, rescaling the pixel size to 0.325 um/px WOULD reproduce ~25 "
            "neighbours almost exactly, and that is precisely why it was rejected: tuning a "
            "measured instrument constant until it reproduces a number from a different "
            "cohort is the silent-corruption failure mode the project brief warns about. "
            "The documented pixel size and the [PAPER] 40 um radius were both left "
            "unchanged.",
            "After excluding 7,357 cells labelled 'dirt' (2.8%, an imaging artefact class) "
            "and images below the minimum cell count, 250,476 cells across 135 images and "
            "35 patients entered CN discovery, spanning 28 cell types.",
            "**The k sweep does not point unambiguously at k=10.** Silhouette peaks at k=9 "
            "(0.319), Davies-Bouldin is minimised at k=11 (1.180), and adjacent-k ARI is "
            "highest at k=17 (0.978) with a strong local peak at k=11 (0.968). The three "
            "criteria disagree, exactly as the method anticipates, which is why all three "
            "are reported rather than one being optimised. k=10 was retained as configured. "
            "Worth noting: the silhouette optimum at k=9 coincides with the nine "
            "neighbourhoods the original authors published, which is mild independent "
            "corroboration of the structure rather than of the specific choice of ten.",
            "**External validation against the authors' own labels (T11).** The source "
            "table ships the published neighbourhood assignments, so the rediscovered CNs "
            "can be scored against them directly with no additional data. Adjusted Rand "
            "index 0.377 and normalised mutual information 0.470 across 250,476 cells. "
            "Every published neighbourhood type is represented among the rediscovered set, "
            "with Bulk tumour recovered at 89.5% overlap, Granulocyte-enriched at 82.2%, a "
            "memory T-cell compartment at 82.0% and Follicle at 72.6%. Agreement is "
            "moderate by construction: the published procedure built each window from the "
            "10 nearest neighbours, whereas CANVAS specifies a fixed 40 um radius followed "
            "by a topic decomposition. This measures whether a CANVAS-style procedure "
            "recovers comparable tissue structure on the same cells, which it does, not "
            "whether the original code was reproduced.",
            "CN names in T2 were read from the marker-enrichment table in this cohort. The "
            "published NSCLC nomenclature was deliberately not reused, and the published "
            "CRC labels were not consulted when naming, so the correspondence in T11 is an "
            "outcome rather than an assumption."],
               figure_ids=["F2", "F3", "F4", "F5"], table_ids=["T2", "T11"]),

        Section("3.2 Label transfer (Stage 2) -- REAL DATA, with a negative result", 2, [
            "The Orion CRC01 specimen was downloaded from the public release "
            "(s3://lin-2023-orion-crc): the registered H&E whole-slide image (981 MB, "
            "78,417 x 57,360 px at 0.325 um/px, 8 pyramid levels) and its single-cell "
            "table (776 MB, 1,620,375 cells, 16 biological markers). Both byte-exact "
            "against the S3 manifest. The 105 GB multiplex image for this specimen was "
            "deliberately not downloaded: the registered H&E and the cell table are what "
            "label transfer needs.",
            "**Coordinate units, again established rather than assumed.** Cell centroids "
            "are pixels in the SAME frame as the H&E: X_centroid maxes at 78,371 against "
            "an image width of 78,417 (ratio 0.999), Y at 55,447 against 57,360 (0.967). "
            "Median cell area of 406 px2 gives a 7.4 um equivalent diameter at 0.325 "
            "um/px, which is a real cell. Same-frame coordinates are what Orion's "
            "same-section design implies, and they make the expected affine the identity.",
            "**Registration verified, not trusted (F6).** 97.1% of sampled cell centroids "
            "fall inside the H&E tissue mask, and an FFT cross-correlation between the "
            "tissue mask and the cell-density map peaks at exactly zero offset. Reported "
            "as a bound rather than a false-precision zero: the residual is below 2.6 um, "
            "the per-pixel resolution of the pyramid level used, comfortably inside the "
            "[PAPER] 5 um threshold. The check is global and rigid and would not detect "
            "local warping.",
            "**Real H&E patches were extracted and are shown in F7**, read as 345 native "
            "px windows (224 px at the 20x working resolution) directly from the "
            "level-0 image through a zarr view, so the 13.5 GB full-resolution array is "
            "never materialised.",
            "**The habitat label transfer itself did NOT succeed, and is reported as a "
            "negative result rather than presented as a working stage.** Transferring the "
            "Schurch CN taxonomy onto Orion requires a shared cell-type vocabulary, and "
            "the two panels do not have one: Schurch resolves 28 phenotypes from 56 "
            "markers, Orion 16 biological markers. Collapsing both onto the 8 lineages "
            "Orion can express produced a degenerate assignment, with one neighbourhood "
            "(CN06) absorbing 64% of all cells at a median cosine similarity of only "
            "0.653. The cause is diagnosable: in the collapsed 8-lineage space CN06 has "
            "the flattest profile of all ten (no lineage above 28.5%), because the "
            "plasma-cell and unassigned populations that define it in the full taxonomy "
            "both fall into catch-all buckets. A featureless centroid is nearest to "
            "everything in cosine distance, so it acts as an attractor.",
            "**F7 doubles as the visual audit that localises the problem.** Grouping real "
            "H&E patches by dominant gated lineage shows the tumour gate is "
            "morphologically correct (malignant glandular epithelium with cribriform "
            "architecture and enlarged hyperchromatic nuclei) and the smooth-muscle gate "
            "likewise, while the T-cell, macrophage and vascular gates reach only 28-46% "
            "dominant-lineage purity and are visibly mixed. The immune gates, not the "
            "tumour gate, are the limiting step.",
            "One earlier suspicion was corrected by this audit. The gated tumour fraction "
            "(57.4%) far exceeds Schurch's (19%), which initially looked like an "
            "over-permissive threshold. The morphology says otherwise: the patches called "
            "tumour are unambiguously tumour. The difference is far more likely to be "
            "sampling, since Orion images a whole tumour-rich resection whereas the "
            "Schurch cores were deliberately positioned at the invasive front to balance "
            "compartments. This is stated as the more probable explanation, not a "
            "settled one.",
            "**Deriving the habitats de novo on Orion resolved it, and is arguably the "
            "more faithful method anyway.** CANVAS discovers CNs from the spatial-omics "
            "modality that is PAIRED with the H&E, which here is Orion. Doing that "
            "removes the vocabulary collapse entirely, because discovery and transfer "
            "then share one panel. Running the identical stage-1 procedure (40 um radius, "
            "topic decomposition, k-means, all via the same tested functions) over the "
            "8 gated lineages on all 1,620,375 cells produced ten habitats each with a "
            "clear identity: pure tumour (CN01), a second tumour compartment (CN04), "
            "myeloid/immune (CN02), smooth-muscle stroma (CN03 and CN09), a B and T cell "
            "lymphoid aggregate (CN05), vasculature (CN08), mixed immune infiltrate "
            "(CN10) and a tumour-adjacent mixed compartment (CN07). No centroid acts as "
            "an attractor.",
            "**These independently derived habitats recapitulate the Schurch taxonomy.** "
            "Matching the two by lineage composition, 8 of 10 Orion habitats align to a "
            "Schurch CN at cosine above 0.89, including bulk tumour at 0.999, smooth "
            "muscle at 0.970 and the lymphoid compartment at 0.971. The two exceptions "
            "are the habitat dominated by unassigned cells (0.559) and one vascular/"
            "stromal pairing (0.725). Two cohorts, two platforms, two independent "
            "derivations, converging structure.",
            "T3 therefore reports real patch counts from the de novo habitats: 13,321 "
            "patches survive the CANVAS purity rules, all ten habitats plus background "
            "represented, the smallest at 87 patches. The rejected cross-platform "
            "transfer gave 2 patches for its smallest class. Only one specimen was "
            "processed, so no train/validation/test split is reported: splitting is a "
            "patient-level operation and one patient cannot be split.",
            "Separately, `scripts/run_stage2_pair.py` was exercised end to end on "
            "synthetic fixtures, which caught and fixed two real bugs: one sample's "
            "failure was killing the whole multi-sample batch, and a split-count "
            "shortfall was discarding every other sample's completed labelling work "
            "rather than saving it."],
               figure_ids=["F6", "F7"], table_ids=["T3"]),

        Section("3.3 H&E model and Method 1 vs Method 2 benchmark (Stage 3)", 2, [
            "The embedding extraction pipeline that Method 2 depends on end to end "
            "(`stage3_model/encoders.py::load_encoder`, `stage3_model/extract_embeddings.py`, "
            "`scripts/run_stage3_extract.py`) is now implemented: openslide tiling, a "
            "simplified CLAM-style tissue mask, colour-based artefact filtering, Macenko "
            "stain normalisation, streaming batch encoding to resumable per-slide parquet "
            "shards. 20 new tests, synthetic images only, zero network calls -- every real "
            "per-backend loader (phikon/resnet50/uni/musk) needs either a download or "
            "gated HF access, so none is exercised in the automated suite.",
            f"The four-way ablation ran on the `--simulate` fixture at {n_seeds_main} seeds "
            "x 4 modes x 15 epochs (F8-F11, T4-T6). mode=none IS Method 1/CANVAS exactly; "
            "k=0 in CANVAS-CTX reduces to the identical head on the identical raw "
            "embedding (Section 4). See Section 3.3.1 for the numbers."],
               figure_ids=["F8", "F9", "F10", "F11"], table_ids=["T4", "T5", "T6"]),

        Section("3.4 WSI inference (Stage 4)", 2, [
            "Not run: needs TCGA-COAD/BRCA slides plus a trained habitat head, neither "
            "available yet. `scripts/run_stage4_infer.py` is written -- DX filtering and "
            "mpp validation via the existing tested `is_diagnostic_slide`/`read_slide_info`, "
            "patch embedding via the Stage 3 pipeline above, and OPTIONAL habitat "
            "prediction + compartment assignment if a checkpoint is supplied. Without one "
            "it stops after embedding caching and says so, rather than fabricating "
            "compartments."], figure_ids=["F12", "F13"]),

        Section("3.5 Spatial features (Stage 5)", 2, [
            "The full 262-feature pipeline ran end to end on 200 simulated habitat maps "
            "(`scripts/run_stage5_features.py --simulate --n-samples 200`): shape "
            "(200, 262), zero all-NaN columns, the runtime assertion on feature count held. "
            "F14 shows the clustered feature matrix and F15 the collinearity structure "
            "`reduce_collinearity` (Louvain, |rho|>0.95) actually finds on this fixture -- "
            "both real computations on synthetic data, not placeholders."],
               figure_ids=["F14", "F15"]),

        Section("3.6 Clinical modelling and the null-calibration check (Stage 6)", 2, [
            "The full chain -- univariate Cox with BH-FDR, consensus clustering, "
            "collinearity reduction, LASSO-Cox selection, RSF importance, multivariable "
            "signature with C-index and time-dependent AUC -- ran against the SAME 262 "
            "simulated features joined to an INDEPENDENTLY RANDOM survival outcome "
            "(exponential event/censoring times, no dependence on any feature), at n=50 "
            "(two seeds) and n=200. This is a calibration check: every stage should come "
            "out null.",
            "It mostly does. Univariate Cox is clean: 0/261 features at q<0.05 (n=200; "
            "one n=50 run showed 2/261, still consistent with noise before FDR). "
            "Collinearity reduction and consensus clustering run correctly and do not see "
            "the outcome at all, so there is nothing to calibrate-check there.",
            f"**The final multivariable signature model is NOT null**, and this is the "
            "single most important finding in this report. Across four independent null "
            "draws the fitted C-index came out 0.80 and 0.75 at n=50 (clinical seeds 42 "
            "and 999), 0.60 at n=200 (seed 42), and "
            f"{_signature_cindex_note(fig)} "
            "Those are separate random draws rather than one controlled series -- the "
            "n=50 and n=200 values differ in clinical seed as well as sample size -- so "
            "read the pattern, not any single number. Every one of them sits above the "
            "0.5 a true null should give. "
            "This is systematic, not a fluke -- confirmed across independent seeds, and "
            "it shrinks as n grows, which is the exact signature of in-sample optimism "
            "bias: `lasso_cox_selection` selects features, `rsf_importance` filters "
            "further, and `fit_signature` fits AND evaluates the multivariable Cox model, "
            "all on the same subjects, with no held-out split anywhere in the chain. With "
            "few samples this cherry-picks noise efficiently; with more samples it is "
            "harder to.",
            "A related, precise discrepancy found while tracing this: `ClinicalConfig."
            "lasso_cv_folds` is declared and documented (\"5-fold internal CV on "
            "C-index\") but is never referenced anywhere else in `signature.py` -- the "
            "actual LASSO step does a single 70/30 split per repeat and takes the "
            "regularisation path's midpoint coefficient, not a cross-validated alpha. Not "
            "fixed here (the file is protected/tested), just flagged: it is a real "
            "doc/code mismatch, and it is plausibly part of why the selection is unstable "
            "at small n.",
            "F21 and F22 need no external data and directly demonstrate the project's two "
            "DELIBERATE deviations from the published method (T10 rows 6-7): the "
            "toroidal-shift interaction null versus label shuffling, and border-corrected "
            "versus uncorrected Ripley's K. Both are reproduced fresh here on synthetic "
            "clustered tissue, not asserted from memory."],
               figure_ids=["F16", "F17", "F18", "F19", "F20", "F21", "F22"],
               table_ids=["T7", "T8", "T9", "T10"]),

        Section("4. Method 1 versus Method 2, and the k=0 equivalence", 1, [
            "mode=none in every figure and table above IS Method 1 (CANVAS): the unchanged "
            "CANVAS head applied directly to one patch's cached embedding, no neighbourhood "
            "context. CANVAS-CTX (Method 2) wraps that same head with a context encoder -- "
            "graph (k-NN deep-set + distance-biased attention), grid2d (local WxW lattice, "
            "2D conv), or grid3d (multi-scale SxWxW cube, 3D conv) -- and at k_neighbours=0 "
            "every one of those context branches degenerates to passing the raw embedding "
            "straight through, which is architecturally identical to mode=none. That "
            "equivalence is enforced by a dedicated test (`tests/test_context_model.py`), "
            "not just asserted in prose, so the k-sweep from 0 upward is a clean ablation: "
            "same encoder, same cached embeddings, same focal loss, same weighted sampler, "
            "same sample-level splits, same head. Any macro-F1 gap from graph/grid2d/grid3d "
            "over none is attributable to spatial context specifically, not to a different "
            "backbone or more parameters overall -- except that grid3d does carry "
            "meaningfully more parameters than grid2d in this run (T5), which is exactly "
            "why PROTOCOL.md calls for a matched-`grid_channels` rerun before trusting a "
            "grid3d-over-grid2d win; that rerun was not performed here (Limitations)."]),

        Section("5. Limitations", 1, [
            "Stated bluntly, in the order that matters most.",
            "**Zero real data.** Every number in this report is either the synthetic "
            "`--simulate` fixture or a null-outcome calibration check. Nothing here says "
            "anything about breast or colorectal tissue.",
            "**The benchmark fixture rewards context by construction.** Oriented bands are "
            "built into `run_final_benchmark.py`'s `simulate()` specifically so spatial "
            "context helps; a large gain there is a machinery check, and the script's own "
            "docstring says so.",
            "**In-sample optimism bias in the signature model** (Section 3.6). Any real "
            "C-index or hazard ratio produced by this same code path on real data needs a "
            "held-out split or bootstrap optimism correction before it goes in a "
            "manuscript.",
            "**Sample sizes, even hypothetically.** The intended real deployment (60 "
            "TCGA-COAD + 60 TCGA-BRCA slides, PAM50-stratified into 5 subtypes) was "
            "already flagged in this project's own config as too small for interpretable "
            "hazard ratios -- `config/pilot.yaml`'s honesty ledger says so explicitly.",
            "**Frozen encoder.** Phikon stays frozen throughout (pilot profile default); "
            "the paper fine-tunes MUSK's final two layers. Expect a lower ceiling on real "
            "data than the published numbers.",
            "**Single-cohort training, cross-cancer application.** The whole pipeline "
            "trains its CN taxonomy and habitat classifier on colorectal tissue (Schurch/"
            "Orion) and applies the trained model to breast unchanged. CANVAS did the "
            "analogous thing (lung to 12 tumour types), but CN biological identity is not "
            "guaranteed to transfer, and T10/the design docs say so.",
            "**Multi-scale pooling is an approximation.** grid3d's `scale_mode=\"pool\"` "
            "(laptop default) average-pools the fine lattice rather than re-encoding at "
            "each scale; the mean of embeddings is not the embedding of the mean under a "
            "non-linear encoder. `scale_mode=\"encode\"` is faithful but triples encoding "
            "cost and was not used here.",
            "**grid3d vs grid2d parameter count not controlled.** grid3d beats grid2d in "
            "T5/F11 with roughly 2.2x the parameters; PROTOCOL.md requires a `grid_channels`"
            "-matched rerun before attributing that gap to the scale axis, which this run "
            "did not do.",
            "**Per-class recall visibility gap.** The main 6-seed benchmark's own "
            "background-command output was piped through `tail`, which (not anticipated "
            "at the time) also truncated the on-disk log, so direct per-seed collapse-"
            "check output only survives for seeds 5-6 of that run. T6/F9's per-class "
            "numbers come from a SEPARATE 3-seed supplementary run instrumented to persist "
            "per-class metrics, not from the main 6-seed run -- both are simulated fixture "
            "runs, but they are not literally the same run.",
            "**Not a git repository.** No commit SHA is available for this run (see "
            "Reproducibility) -- there is no version control on this project directory."]),

        Section("6. Reproducibility", 1, [
            f"Config: `{stats['config_path']}`, resolved through its `inherit:` chain "
            f"(`config/pilot.yaml` -> `config/default.yaml`) by "
            f"`canvas_brca.utils.config.load_config` (new this run -- no prior script "
            f"resolved that chain). SHA-256 of the merged, JSON-serialised config: "
            f"`{stats['config_hash_sha256']}...`. Project seed: {stats['project_seed']}. "
            f"Benchmark seeds: 1-6 (main run), 1-3 (per-class supplementary run). "
            f"Null-calibration seeds: 42 and 999 (n=50), 42 (n=200).",
            f"Git SHA: **{stats['git_sha']}**.",
            "**Software environment.** Every third-party library the pipeline "
            "uses is listed below with its installed version and the role it "
            "plays, so the environment can be reconstructed from this report "
            "without reading the source. Python "
            f"{stats['python_version']} on {stats['platform']}.",
            _software_table(stats["versions"]),
            "**Optional dependencies and the fallbacks taken.** Each of these "
            "is absent from this run, and each has a documented fallback rather "
            "than a silent skip:\n\n"
            + "\n".join(f"- **{name}** — {role}" for name, role in OPTIONAL_SOFTWARE),
            "Approximate runtime this session (CPU-only, Ryzen 7000, no GPU): stage5 "
            "`--simulate --n-samples 200` ~17.5 min; 6-seed x 4-mode x 15-epoch benchmark "
            "~85 min (grid3d dominates -- pilot.yaml's own comment already warns it is "
            "~4x grid2d); 3-seed supplementary per-class/confusion-matrix run ~45 min; "
            "null-calibration chain (n=200) ~2 min; report figure/table generation "
            "(this file) ~5-10 min, dominated by F18's consensus-clustering k-sweep.",
            "Exact commands to regenerate every artefact referenced in this report:",
            "```\n"
            "python -m pytest tests/ -q\n"
            "python scripts/run_stage5_features.py --simulate --n-samples 200\n"
            "python scripts/run_final_benchmark.py --simulate --modes none graph grid2d "
            "grid3d --seeds 1 2 3 4 5 6 --epochs 15 --window 7\n"
            "python scripts/run_final_benchmark.py --simulate --modes none graph grid2d "
            "grid3d --seeds 1 2 3 --epochs 15 --window 7 --outdir results/per_class_run\n"
            "python scripts/validate_stage6_null.py --features data/interim/sim_features.parquet "
            "--endpoint OS --seed 42 --outdir results/null_check_n200\n"
            "python scripts/run_report.py\n"
            "```"]),
    ]
    return sections


def _software_table(versions: dict[str, str]) -> str:
    """Markdown table of every library, its version and what it is used for."""
    rows = ["| library | version | import | role |", "|---|---|---|---|"]
    for dist, imp, role in SOFTWARE:
        rows.append(f"| `{dist}` | {versions.get(dist, 'not installed')} | "
                    f"`{imp}` | {role} |")
    return "\n".join(rows)


def _signature_cindex_note(fig: dict) -> str:
    stats = fig.get("F20", {}).get("stats")
    if stats and stats.get("c_index") is not None:
        return f"{stats['c_index']:.3f} at n={stats['n']} (F20)."
    return "an undetermined value at n=200 -- F20 could not fit a model on this run (F20)."


# ==================================================================== markdown


def _df_to_md(df: pd.DataFrame, max_rows: int = 25) -> str:
    if len(df) > max_rows:
        shown = df.head(max_rows)
        note = f"\n\n*(showing {max_rows} of {len(df)} rows; full table in the CSV)*"
    else:
        shown, note = df, ""
    return shown.to_markdown(index=False, floatfmt=".4g") + note


def render_markdown(sections: list[Section], fig: dict, tab: dict, figures_dir: str) -> str:
    lines = ["# CANVAS-BRCA / CANVAS-CTX: Analysis Report", ""]
    for sec in sections:
        lines.append(f"{'#' * (sec.level + 1)} {sec.heading}")
        lines.append("")
        for p in sec.paragraphs:
            lines.append(p)
            lines.append("")
        for fid in sec.figure_ids:
            f = fig[fid]
            rel = Path(f["paths"]["png"]).name
            lines.append(f"**{f['id']}. {f['title']}** -- *{f['source']}*")
            lines.append("")
            lines.append(f"![{f['id']}]({figures_dir}/{rel})")
            lines.append("")
            lines.append(f["caption"])
            lines.append("")
        for tid in sec.table_ids:
            t = tab[tid]
            lines.append(f"**{t['id']}. {t['title']}** -- *{t['source']}*")
            lines.append("")
            lines.append(_df_to_md(t["df"]))
            lines.append("")
            lines.append(t["caption"])
            lines.append("")
    return "\n".join(lines)


# ======================================================================= html


_HTML_HEAD = f"""<!doctype html><html><head><meta charset="utf-8">
<title>CANVAS-BRCA / CANVAS-CTX Analysis Report</title>
<style>
body {{ font-family: Calibri, 'Segoe UI', sans-serif; max-width: 960px; margin: 2rem auto;
       padding: 0 1.5rem; line-height: 1.55; color: #1a1a1a; }}
h1 {{ color: {NAVY}; border-bottom: 3px solid {NAVY}; padding-bottom: .3rem; }}
h2 {{ color: {NAVY}; border-bottom: 1px solid #ccc; padding-bottom: .2rem; margin-top: 2.2rem; }}
h3 {{ color: {STEEL_BLUE}; margin-top: 1.6rem; }}
.source-real {{ color: #006400; font-weight: 600; }}
.source-sim {{ color: {STEEL_BLUE}; font-weight: 600; }}
.source-missing {{ color: #b00020; font-weight: 600; }}
.caption {{ font-size: 0.92em; color: #444; margin: .3rem 0 1.2rem 0; }}
img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 4px; }}
table {{ border-collapse: collapse; margin: .8rem 0; font-size: .88em; max-width: 100%; display: block;
        overflow-x: auto; }}
th, td {{ border: 1px solid #ccc; padding: 4px 8px; text-align: left; }}
th {{ background: {NAVY}; color: white; }}
tr:nth-child(even) {{ background: #f4f7fa; }}
code, pre {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
pre {{ padding: 10px; overflow-x: auto; }}
</style></head><body>
<h1>CANVAS-BRCA / CANVAS-CTX: Analysis Report</h1>
"""


def _md_inline(text: str) -> str:
    """Minimal inline markdown: **bold** and `code`."""
    import re
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    return re.sub(r"`(.+?)`", r"<code>\1</code>", text)


def _md_table_to_html(md: str) -> str:
    """Render a pipe table written in a prose paragraph as a real HTML table."""
    lines = [ln.strip() for ln in md.strip().splitlines() if ln.strip()]
    rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines]
    if len(rows) < 2:
        return f"<p>{md}</p>"
    header, body = rows[0], rows[2:]          # rows[1] is the --- separator
    out = ["<table><thead><tr>"]
    out += [f"<th>{_md_inline(c)}</th>" for c in header]
    out.append("</tr></thead><tbody>")
    for r in body:
        out.append("<tr>" + "".join(f"<td>{_md_inline(c)}</td>" for c in r) + "</tr>")
    out.append("</tbody></table>")
    return "".join(out)


def _md_bullets_to_html(md: str) -> str:
    lead, items = [], []
    for ln in md.splitlines():
        (items if ln.lstrip().startswith("- ") else lead).append(ln)
    html = f"<p>{_md_inline(' '.join(x for x in lead if x.strip()))}</p>" if any(
        x.strip() for x in lead) else ""
    html += "<ul>" + "".join(
        f"<li>{_md_inline(i.lstrip()[2:])}</li>" for i in items) + "</ul>"
    return html


def _source_class(source: str) -> str:
    if source.startswith("REAL"):
        return "source-real"
    if source.startswith("SIMULATED"):
        return "source-sim"
    if source.startswith("MISSING"):
        return "source-missing"
    return ""


def _b64_img(png_path: str) -> str:
    data = base64.b64encode(Path(png_path).read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


def render_html(sections: list[Section], fig: dict, tab: dict) -> str:
    parts = [_HTML_HEAD]
    for sec in sections:
        tag = "h2" if sec.level == 1 else "h3"
        parts.append(f"<{tag}>{sec.heading}</{tag}>")
        for p in sec.paragraphs:
            if p.startswith("```"):
                parts.append(f"<pre>{p.strip('`').lstrip(chr(10))}</pre>")
            elif p.lstrip().startswith("| "):
                parts.append(_md_table_to_html(p))
            elif p.lstrip().startswith("- ") or "\n- " in p:
                parts.append(_md_bullets_to_html(p))
            else:
                parts.append(f"<p>{p}</p>")
        for fid in sec.figure_ids:
            f = fig[fid]
            parts.append(f"<h4>{f['id']}. {f['title']} "
                        f"<span class='{_source_class(f['source'])}'>[{f['source']}]</span></h4>")
            parts.append(f"<img src='{_b64_img(f['paths']['png'])}' alt='{f['id']}'>")
            parts.append(f"<div class='caption'>{f['caption']}</div>")
        for tid in sec.table_ids:
            t = tab[tid]
            parts.append(f"<h4>{t['id']}. {t['title']} "
                        f"<span class='{_source_class(t['source'])}'>[{t['source']}]</span></h4>")
            shown = t["df"].head(25)
            parts.append(shown.to_html(index=False, float_format=lambda x: f"{x:.4g}"))
            if len(t["df"]) > 25:
                parts.append(f"<div class='caption'>(showing 25 of {len(t['df'])} rows; "
                            f"full table at {t['csv_path']})</div>")
            parts.append(f"<div class='caption'>{t['caption']}</div>")
    parts.append("</body></html>")
    return "\n".join(parts)


# ======================================================================= docx


def render_docx(sections: list[Section], fig: dict, tab: dict, out_path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    navy = RGBColor(0x1C, 0x2B, 0x4A)
    steel = RGBColor(0x24, 0x71, 0xA3)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("CANVAS-BRCA / CANVAS-CTX: Analysis Report", level=0)
    for run in title.runs:
        run.font.color.rgb = navy

    for sec in sections:
        h = doc.add_heading(sec.heading, level=min(sec.level + 1, 4))
        for run in h.runs:
            run.font.color.rgb = navy if sec.level == 1 else steel
        for p_text in sec.paragraphs:
            if p_text.startswith("```"):
                code = p_text.strip("`").lstrip("\n")
                p = doc.add_paragraph(code)
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                continue
            if p_text.lstrip().startswith("| "):
                # a pipe table written in prose becomes a real Word table
                lines = [ln.strip() for ln in p_text.strip().splitlines() if ln.strip()]
                rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in lines]
                if len(rows) >= 2:
                    header, body = rows[0], rows[2:]
                    tbl = doc.add_table(rows=1, cols=len(header))
                    tbl.style = "Light Grid Accent 1"
                    for i, c in enumerate(header):
                        tbl.rows[0].cells[i].text = c.replace("`", "")
                    for r in body:
                        cells = tbl.add_row().cells
                        for i, c in enumerate(r[:len(header)]):
                            cells[i].text = c.replace("`", "").replace("**", "")
                    doc.add_paragraph()
                    continue
            p = doc.add_paragraph(p_text.replace("**", ""))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for fid in sec.figure_ids:
            f = fig[fid]
            h4 = doc.add_heading(f"{f['id']}. {f['title']}  [{f['source']}]", level=4)
            for run in h4.runs:
                run.font.color.rgb = steel
            doc.add_picture(f["paths"]["png"], width=Cm(15))
            cap = doc.add_paragraph(f["caption"])
            cap.italic = True
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)

        for tid in sec.table_ids:
            t = tab[tid]
            h4 = doc.add_heading(f"{t['id']}. {t['title']}  [{t['source']}]", level=4)
            for run in h4.runs:
                run.font.color.rgb = steel
            shown = t["df"].head(20)
            table = doc.add_table(rows=1, cols=len(shown.columns))
            table.style = "Light Grid Accent 1"
            for i, col in enumerate(shown.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in shown.iterrows():
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val:.4g}" if isinstance(val, float) else str(val)
            if len(t["df"]) > 20:
                doc.add_paragraph(f"(showing 20 of {len(t['df'])} rows; full table at "
                                 f"{t['csv_path']})").italic = True
            cap = doc.add_paragraph(t["caption"])
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ==================================================================== entrypoint


def build_report(figures_dir: str = "figures", tables_dir: str = "results/tables",
                 reports_dir: str = "reports",
                 config_path: str = "config/crc_train_brca_apply.yaml") -> dict:
    fig, tab = _run_all(figures_dir, tables_dir)
    stats = _reproducibility_stats(config_path)
    sections = build_sections(fig, tab, stats)

    # Completeness contract: EVERY generated figure and table must actually be
    # placed in a section, or it silently never reaches the report. Fail loudly
    # rather than quietly shipping an incomplete document.
    placed_figs = {fid for s in sections for fid in s.figure_ids}
    placed_tabs = {tid for s in sections for tid in s.table_ids}
    orphan_figs = sorted(set(fig) - placed_figs, key=lambda s: int(s[1:]))
    orphan_tabs = sorted(set(tab) - placed_tabs, key=lambda s: int(s[1:]))
    if orphan_figs or orphan_tabs:
        raise AssertionError(
            f"report would omit generated content: figures {orphan_figs}, "
            f"tables {orphan_tabs}. Add them to a Section in build_sections()."
        )

    Path(reports_dir).mkdir(parents=True, exist_ok=True)
    md = render_markdown(sections, fig, tab, figures_dir)
    (Path(reports_dir) / "analysis_report.md").write_text(md, encoding="utf-8")

    html = render_html(sections, fig, tab)
    (Path(reports_dir) / "analysis_report.html").write_text(html, encoding="utf-8")

    render_docx(sections, fig, tab, str(Path(reports_dir) / "analysis_report.docx"))

    n_missing_fig = sum(1 for f in fig.values() if f["source"] == "MISSING DATA")
    n_missing_tab = sum(1 for t in tab.values() if t["source"] == "MISSING DATA")
    return {"figures": fig, "tables": tab, "stats": stats,
           "n_missing_figures": n_missing_fig, "n_missing_tables": n_missing_tab,
           "reports_dir": reports_dir}
