"""Manuscript draft generator: reports/manuscript.{md,docx}.

THE HARD RULE, same as report.py: no fabricated results and no fabricated
citations. Sections that can be written now (Introduction, Methods,
Limitations, figure legends, references) are written in full. Every slot
that requires data this project does not have is marked

    [RESULTS PENDING - requires <exact file>; unblocks <figure/table ids>]

and is deliberately left empty rather than filled with a plausible number.
A reviewer, a supervisor, or a future you must be able to see at a glance
which claims are supported and which are still holes.

Every reference in REFERENCES was verified against PubMed/publisher pages
rather than recalled. `verified` records that; anything that could not be
independently confirmed says so in `note` and is flagged in the output.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from .style import NAVY, STEEL_BLUE


@dataclass
class Ref:
    key: str
    citation: str
    pmid: str | None = None
    doi: str | None = None
    verified: bool = True
    note: str = ""


REFERENCES: list[Ref] = [
    Ref("canvas", "Li et al. Virtual spatial profiling of tumour ecological habitats "
        "from routine histology. Cell (2026).",
        pmid=None, doi="10.1016/j.cell.2026.05.031", verified=False,
        note="DOI as supplied in the project brief; could not be independently "
             "confirmed against PubMed at the time of writing. VERIFY BEFORE SUBMISSION."),
    Ref("schurch", "Schurch CM, Bhate SS, Barlow GL, et al. Coordinated cellular "
        "neighborhoods orchestrate antitumoral immunity at the colorectal cancer "
        "invasive front. Cell. 2020;182(5):1341-1359.e19.",
        pmid="32763154", doi="10.1016/j.cell.2020.10.021"),
    Ref("orion", "Lin JR, Chen YA, Campton D, et al. High-plex immunofluorescence "
        "imaging and traditional histology of the same tissue section for discovering "
        "image-based biomarkers. Nat Cancer. 2023;4(7):1036-1052.",
        pmid=None, doi="10.1038/s43018-023-00576-1",
        note="DOI verified against the Nature Cancer article page; PMID not "
             "captured during verification."),
    Ref("danenberg", "Danenberg E, Bardwell H, Zanotelli VRT, et al. Breast tumor "
        "microenvironment structures are associated with genomic features and clinical "
        "outcome. Nat Genet. 2022;54(5):660-669.",
        pmid="35437329", doi="10.1038/s41588-022-01041-y"),
    Ref("jackson", "Jackson HW, Fischer JR, Zanotelli VRT, et al. The single-cell "
        "pathology landscape of breast cancer. Nature. 2020;578(7796):615-620.",
        pmid="31959985", doi="10.1038/s41586-019-1876-x"),
    Ref("keren", "Keren L, Bosse M, Marquez D, et al. A structured tumor-immune "
        "microenvironment in triple negative breast cancer revealed by multiplexed ion "
        "beam imaging. Cell. 2018;174(6):1373-1387.e19.",
        pmid="30193111", doi="10.1016/j.cell.2018.08.039"),
    Ref("tcgacdr", "Liu J, Lichtenberg T, Hoadley KA, et al. An integrated TCGA "
        "pan-cancer clinical data resource to drive high-quality survival outcome "
        "analytics. Cell. 2018;173(2):400-416.e11.",
        pmid="29625055", doi="10.1016/j.cell.2018.02.052"),
    Ref("musk", "Xiang J, Wang X, Zhang X, et al. A vision-language foundation model "
        "for precision oncology. Nature. 2025;638(8051):769-778.",
        pmid=None, doi="10.1038/s41586-024-08378-w",
        note="Volume/pages verified against the Nature article page."),
    Ref("uni", "Chen RJ, Ding T, Lu MY, et al. Towards a general-purpose foundation "
        "model for computational pathology. Nat Med. 2024;30(3):850-862.",
        pmid=None, doi="10.1038/s41591-024-02857-3"),
    Ref("phikon", "Filiot A, Ghermi R, Olivier A, et al. Scaling self-supervised "
        "learning for histopathology with masked image modeling. medRxiv (2023).",
        pmid=None, doi="10.1101/2023.07.21.23292757",
        note="Preprint. Check for a peer-reviewed version before submission."),
    Ref("clam", "Lu MY, Williamson DFK, Chen TY, et al. Data-efficient and weakly "
        "supervised computational pathology on whole-slide images. Nat Biomed Eng. "
        "2021;5(6):555-570.", pmid="33649564", doi="10.1038/s41551-020-00682-w"),
    Ref("titan", "Ding T, Wagner SJ, Song AH, et al. A multimodal whole-slide "
        "foundation model for pathology. Nat Med. 2025;31(11).",
        pmid=None, doi="10.1038/s41591-025-03982-3",
        note="Also available as arXiv:2411.19666."),
    Ref("macenko", "Macenko M, Niethammer M, Marron JS, et al. A method for "
        "normalizing histology slides for quantitative analysis. IEEE ISBI. "
        "2009:1107-1110.", pmid=None, doi="10.1109/ISBI.2009.5193250",
        note="Conference paper; no PMID."),
]

PENDING = ("[RESULTS PENDING - requires {data}; unblocks {unblocks}. "
          "Do not draft this paragraph until the data exists.]")


@dataclass
class MSection:
    heading: str
    level: int
    paragraphs: list[str] = field(default_factory=list)


def _pending(data: str, unblocks: str) -> str:
    return PENDING.format(data=data, unblocks=unblocks)


def build_manuscript_sections() -> list[MSection]:
    return [
        MSection("Title", 1, [
            "**Spatial context improves inference of tumour ecological habitats from "
            "routine histology: a colorectal-trained, breast-applied re-implementation "
            "of CANVAS**",
            "*Author list, affiliations and corresponding author to be completed.*",
        ]),

        MSection("Abstract", 1, [
            "**Background.** Tumour ecological habitats, defined from spatial "
            "proteomics as recurrent cellular neighbourhoods (CNs), can be inferred "
            "from haematoxylin and eosin (H&E) histology alone [ref: canvas]. The "
            "published approach defines each habitat label from a 40 um cellular "
            "neighbourhood but predicts it from a single 224x224 patch viewed in "
            "isolation, so the label is contextual while the predictor is not.",
            "**Methods.** We re-implemented the CANVAS pipeline end to end and "
            "extended it with CANVAS-CTX, which predicts the habitat of a patch from "
            "that patch plus its spatial neighbours through three interchangeable "
            "context encoders (k-nearest-neighbour deep set, two-dimensional feature "
            "grid, multi-scale three-dimensional grid) feeding the unchanged CANVAS "
            "classification head. Because k = 0 reduces CANVAS-CTX exactly to the "
            "published per-patch model, the neighbour sweep is itself a controlled "
            "ablation. Owing to the scarcity of paired same-section spatial proteomics "
            "and H&E in breast tissue, CNs and the habitat classifier were derived in "
            "colorectal cancer [ref: schurch; ref: orion] and applied unchanged to "
            "breast, mirroring the cross-tumour-type transfer reported in the original "
            "work.",
            "**Results.** Applying the neighbourhood-discovery procedure to 250,476 "
            "single cells profiled by 56-plex CODEX across 135 colorectal regions from "
            "35 patients recovered ten spatially coherent cellular neighbourhoods "
            "spanning tumour, boundary, lymphoid, myeloid, stromal and vascular "
            "compartments. Validated against the neighbourhood labels published with "
            "the source cohort, agreement was moderate and structurally complete "
            "(adjusted Rand index 0.377; every published neighbourhood type "
            "represented; bulk tumour recovered at 89.5 percent overlap). " +
            _pending("Orion paired IF/H&E specimens, TCGA-COAD and TCGA-BRCA "
                    "diagnostic slides",
                    "the habitat-inference, transfer and prognostic results"),
            "**Conclusions.** " + _pending(
                "all of the above", "Abstract conclusions"),
        ]),

        MSection("Introduction", 1, [
            "The spatial organisation of the tumour microenvironment carries "
            "prognostic information that is invisible to bulk molecular assays. "
            "Multiplexed imaging platforms have established that tumours are built "
            "from recurrent, discrete cellular neighbourhoods rather than from "
            "randomly mixed cell populations: co-detection by indexing (CODEX) "
            "identified conserved neighbourhoods at the colorectal cancer invasive "
            "front whose organisation stratified survival [ref: schurch]; imaging mass "
            "cytometry defined analogous structures across 693 breast tumours "
            "[ref: danenberg] and a single-cell pathology landscape across 352 "
            "further cases [ref: jackson]; and multiplexed ion beam imaging showed "
            "that the degree of immune compartmentalisation in triple-negative breast "
            "cancer is itself prognostic [ref: keren].",
            "These platforms are, however, expensive, low-throughput and absent from "
            "routine practice, whereas H&E histology is universal. Recent work "
            "demonstrated that habitat identity defined from spatial proteomics can be "
            "transferred onto co-registered H&E and predicted from morphology alone "
            "using a pathology foundation model, enabling spatial habitat profiling of "
            "large archival cohorts [ref: canvas]. This depends on pathology "
            "foundation models that have advanced rapidly, from self-supervised "
            "encoders trained on pan-cancer tiles [ref: phikon; ref: uni] to "
            "vision-language models trained on paired image and text corpora "
            "[ref: musk].",
            "A methodological gap remains. The habitat label assigned to a patch is "
            "defined by the composition of a 40 um neighbourhood surrounding it, yet "
            "the classifier that must reproduce that label sees only the patch itself. "
            "Information that defines the target is withheld from the predictor. The "
            "wider field has moved toward spatial aggregation for precisely this "
            "reason, arranging patch features on a two-dimensional lattice so that "
            "positional encoding and relative distance are preserved [ref: titan]. "
            "Directionality matters biologically: invasive fronts, duct walls and "
            "tumour-stroma interfaces are oriented structures, and a permutation-"
            "invariant summary of what is nearby cannot represent which direction it "
            "lies in.",
            "We therefore re-implemented the published pipeline faithfully and added "
            "CANVAS-CTX, a context-aware variant in which the unchanged classification "
            "head receives, in addition to the index patch embedding, a representation "
            "of its spatial neighbourhood produced by one of three encoders. Setting "
            "the neighbour count to zero recovers the published model exactly, so any "
            "difference is attributable to spatial context rather than to a different "
            "backbone, a different loss, or additional capacity in the head.",
            "A second constraint shaped the design. The pipeline requires spatial "
            "proteomics and H&E from the same tissue section. In breast cancer this "
            "pairing is scarce, and the available resource documents its own image "
            "alignment as suitable for regional annotation rather than sub-cellular "
            "correspondence. Colorectal cancer, by contrast, offers same-section "
            "18-plex immunofluorescence with matched H&E across 74 resections "
            "[ref: orion], alongside the CODEX cohort in which cellular neighbourhoods "
            "were originally defined [ref: schurch]. We consequently derive habitats "
            "and train the classifier in colorectal tissue and apply the trained model "
            "unchanged to breast, treating breast as the generalisation test rather "
            "than the foundation. The original study performed the structurally "
            "equivalent experiment, applying a lung-trained model across multiple "
            "tumour types [ref: canvas].",
        ]),

        MSection("Results", 1, []),

        MSection("Cellular neighbourhoods in colorectal cancer", 2, [
            "We analysed 258,385 single cells profiled by 56-plex CODEX across 140 "
            "imaged regions from 35 patients with colorectal carcinoma "
            "[ref: schurch]. Coordinates in the source table are recorded in image "
            "pixels on a 1920x1440 grid and were converted to microns at the 377.44 "
            "nm/pixel lateral resolution documented for this collection, giving a "
            "724 x 543 um field of view per region. The conversion was verified "
            "against the data itself: median nearest-neighbour centroid spacing after "
            "conversion is 6.8 um, consistent with densely packed tissue. After "
            "excluding 7,357 cells (2.8 percent) assigned to an imaging-artefact class "
            "and regions falling below the minimum cell count, 250,476 cells across 135 "
            "regions and 28 cell types entered neighbourhood discovery.",
            "At the 40 um radius specified by the method, neighbourhoods contained a "
            "median of 33 cells, above the approximately 25 reported for the lung "
            "cohort in which the method was developed. This reflects tissue density "
            "rather than a scaling error: the colorectal invasive front in this cohort "
            "carries a median of 4,890 cells per square millimetre. We note explicitly "
            "that adopting the 0.325 um/pixel resolution quoted for other installations "
            "of this imaging platform would reproduce approximately 25 neighbours, and "
            "that we rejected this adjustment because calibrating a documented "
            "instrument constant to reproduce a value observed in a different tissue "
            "would confound density with scale. The radius and the documented pixel "
            "size were both retained unchanged.",
            "The three selection criteria did not converge on a single solution. "
            "Silhouette width was maximised at nine neighbourhoods (0.319), the "
            "Davies-Bouldin index minimised at eleven (1.180), and the adjusted Rand "
            "index between adjacent partitions peaked at seventeen (0.978) with a "
            "further strong local maximum at eleven (0.968). We report all three rather "
            "than optimising any one, and retained ten neighbourhoods as specified by "
            "the protocol. The silhouette optimum at nine coincides with the nine "
            "neighbourhoods reported by the original investigators of this cohort.",
            "Neighbourhoods were named from cell-type enrichment within this cohort "
            "(Figure 1C, Supplementary Table S2) and comprise bulk tumour (12.1 percent "
            "of cells), a tumour boundary myeloid compartment (9.5 percent), "
            "cytotoxic-infiltrated stroma (11.7 percent), lymphoid follicle (6.4 "
            "percent), pan-immune cytotoxic (11.1 percent), memory CD4 T-cell (5.0 "
            "percent), plasma-cell rich (7.9 percent), macrophage-adipose (16.6 "
            "percent), smooth muscle-lymphatic (12.9 percent) and granulocyte-enriched "
            "(6.7 percent) structures. Plotting cells at their measured coordinates confirms that "
            "these are spatially coherent tissue domains rather than dispersed label "
            "assignments (Figure 1D).",
            "Because the source table also carries the neighbourhood assignments "
            "published by the original investigators, the rediscovered neighbourhoods "
            "could be validated externally without additional data. Agreement across "
            "250,476 cells was moderate (adjusted Rand index 0.377, normalised mutual "
            "information 0.470), with every published neighbourhood type represented "
            "among the rediscovered set: bulk tumour was recovered at 89.5 percent "
            "overlap, granulocyte-enriched at 82.2 percent, a memory T-cell compartment "
            "at 82.0 percent and lymphoid follicle at 72.6 percent (Supplementary Table "
            "S11). Partial rather than complete agreement is expected by construction, "
            "since the published analysis defined each window from the ten nearest "
            "neighbours whereas the present method applies a fixed 40 um radius "
            "followed by a topic decomposition. The comparison therefore establishes "
            "that a CANVAS-style procedure recovers comparable tissue architecture from "
            "the same cells, not that the original implementation was reproduced.",
        ]),

        MSection("Habitat label transfer onto matched H&E", 2, [
            _pending("8-12 Orion CRC specimens (s3://lin-2023-orion-crc/data): "
                    "single-cell tables, registered H&E, segmentation masks",
                    "Figures 2A-2C, Table S3"),
            "*Planned content: median registration residual in microns against the 5 um "
            "acceptance threshold, the proportion of H&E nuclei matched within that "
            "threshold, patch counts retained by the purity rules, and the class "
            "distribution across the patient-level train, validation and test split.*",
        ]),

        MSection("Spatial context improves habitat inference from H&E", 2, [
            _pending("cached patch embeddings from the Orion cohort with transferred "
                    "habitat labels", "Figures 3A-3D, Tables 1 and S6"),
            "*Planned content: macro-F1 and Cohen's kappa with bootstrap confidence "
            "intervals for each of the four context modes across at least six seeds, "
            "the paired Wilcoxon test of each context mode against the per-patch "
            "baseline, per-class recall for every habitat, and the parameter-matched "
            "comparison between the two-dimensional and multi-scale grids. Note that a "
            "paired signed-rank test over six seeds cannot yield p below 0.0312; report "
            "effect size with dispersion, not the p value alone.*",
            "*Machinery validation performed to date, on synthetic data only and not "
            "for inclusion as a finding: the four-mode ablation executes correctly "
            "across six seeds with no class collapse, confirming that the training "
            "loop, the sample-level split and the paired comparison behave as "
            "specified. The synthetic fixture is constructed with oriented bands that "
            "reward spatial context by design and therefore carries no biological "
            "information whatsoever.*",
        ]),

        MSection("Habitat composition of colorectal and breast cohorts", 2, [
            _pending("60 TCGA-COAD and 60 TCGA-BRCA diagnostic (DX) slides",
                    "Figures 4A-4C, Table S4"),
            "*Planned content: whole-slide habitat maps, the tumour bulk and leading "
            "edge partition, and quantified habitat composition shift between the "
            "training cancer type and the transfer cancer type. Any habitat that is "
            "essentially absent in breast should be reported explicitly, as it would "
            "indicate a colorectal neighbourhood with no breast counterpart.*",
        ]),

        MSection("Spatial features and prognostic association", 2, [
            _pending("habitat maps from the preceding section joined to TCGA-CDR "
                    "outcomes", "Figures 5A-5D, Tables 2, S7 and S9"),
            "*Planned content: the 262-feature matrix, univariate Cox hazard ratios "
            "with 95 percent confidence intervals and Benjamini-Hochberg q values per "
            "habitat and compartment, ecotypes from consensus clustering, and the "
            "multivariable signature. Report the number of EVENTS in every stratum. Do "
            "not report a hazard ratio for any stratum containing fewer than ten "
            "events.*",
            "*Critical methodological caution established during pipeline validation: "
            "when the selection chain (LASSO-Cox, then random survival forest "
            "permutation importance, then multivariable Cox) is fitted and evaluated on "
            "the same subjects, it returns concordance indices between 0.60 and 0.80 on "
            "outcomes that are pure noise, with the inflation shrinking as sample size "
            "grows. Any concordance index reported here must therefore come from a "
            "held-out split or carry an explicit bootstrap optimism correction.*",
        ]),

        MSection("Discussion", 1, [
            _pending("all Results sections", "Discussion"),
            "*Planned structure. First paragraph: restate the principal finding, "
            "namely whether spatial context improves habitat inference over the "
            "per-patch baseline and by how much, with dispersion across seeds. Second: "
            "interpret the direction and magnitude of the colorectal-to-breast habitat "
            "shift and what it implies about the biological transferability of "
            "colorectal neighbourhoods. Third: relate the observed effect to spatially "
            "aware architectures elsewhere in the field [ref: titan] and state what the "
            "attention maps show that a per-patch model cannot produce. Fourth: state "
            "the clinical implication honestly, which at pilot scale is that the "
            "pipeline is feasible rather than that any effect size is established.*",
        ]),

        MSection("Limitations", 1, [
            "The cellular neighbourhoods underpinning every downstream analysis are "
            "derived from colorectal tissue and applied to breast without "
            "recalibration. Colorectal and breast neighbourhoods are not biologically "
            "identical, and habitat composition shift between the two cohorts is "
            "reported rather than minimised or concealed.",
            "The habitat classifier is trained on a frozen encoder. The published "
            "method fine-tunes the final two layers of its backbone; freezing is "
            "required here for CPU feasibility and, more importantly, to keep the "
            "comparison between the per-patch and context-aware models fair, since the "
            "latter necessarily trains on cached embeddings. A frozen comparison is "
            "internally valid but establishes a lower ceiling than the published "
            "configuration.",
            "The multi-scale grid fills its scale axis by average-pooling the fine "
            "lattice rather than by re-encoding each magnification. Because the encoder "
            "is non-linear, the mean of embeddings is not the embedding of the mean, so "
            "the coarse scales are approximations. Faithful re-encoding triples "
            "encoding cost and was not performed.",
            "Deployment cohorts of 60 slides per cancer type, stratified across five "
            "intrinsic subtypes, are adequate to demonstrate that the pipeline runs and "
            "inadequate to estimate effect sizes. Hazard ratios from strata containing "
            "few events are not interpretable and are not reported as findings.",
            "Interaction permutation counts are reduced relative to the published "
            "protocol for tractability, which widens the uncertainty on weak "
            "habitat-pair associations while leaving strong ones stable.",
            "Two deliberate departures from the published statistical procedure are "
            "documented in Methods and quantified in the supplementary material: a "
            "toroidal-shift null in place of label shuffling, and edge-corrected "
            "spatial dispersion estimators. Both are more conservative than the "
            "procedures they replace, and both change reported significance.",
        ]),

        MSection("Methods", 1, [
            "*The parameters below marked [PAPER] follow the published protocol "
            "[ref: canvas] and were not altered. Departures are stated explicitly, "
            "with rationale, and are itemised in full in Supplementary Table S10.*",
        ]),

        MSection("Cellular neighbourhood discovery", 2, [
            "Single-cell tables with spatial coordinates and cell-type assignments "
            "were used to construct, for each index cell, the local neighbourhood of "
            "all cells within a 40 um radius [PAPER], which yields approximately 25 "
            "neighbours. Coordinate units were verified before analysis; the "
            "implementation refuses to proceed when the coordinate span is implausible "
            "for microns, because applying a micron radius to pixel coordinates "
            "silently collapses every neighbourhood while still producing clean-looking "
            "clusters. Neighbourhood composition vectors were decomposed by "
            "spatial latent Dirichlet allocation and the resulting topic proportions "
            "clustered by k-means. The number of clusters was swept from 5 to 20 "
            "[PAPER] and scored by silhouette width, Davies-Bouldin index and the "
            "adjusted Rand index between adjacent partitions; all three are reported "
            "and none is optimised alone.",
            "**Departure.** Neighbourhoods were derived from a 56-plex colorectal "
            "CODEX cohort [ref: schurch] rather than the 41-plex lung cohort of the "
            "original study, for the paired-data reasons given in the Introduction.",
        ]),

        MSection("Habitat label transfer", 2, [
            "Spatial-omics cell centroids were mapped into H&E pixel space by an "
            "affine transform. Each H&E-segmented nucleus was assigned the CN of its "
            "nearest transformed spatial cell, subject to a centroid-to-centroid "
            "threshold of 5 um [PAPER]; unmatched nuclei were discarded. Registration "
            "residuals were measured and reported before label transfer proceeded, "
            "rather than assumed from the pre-registered status of the source data.",
            "Patches of 224x224 pixels [PAPER] were labelled by the published purity "
            "rules, applied exactly: five or fewer annotated cells assigns the "
            "background class; more than fifteen annotated cells with a dominant CN "
            "comprising at least 60 percent of local composition assigns that dominant "
            "CN; all intermediate cases are discarded [PAPER]. Train, validation and "
            "test partitions were formed at the patient level and never at the patch "
            "level [PAPER]. This is not a stylistic preference: patch-level splitting "
            "combined with spatial context places neighbouring patches on both sides of "
            "the split and renders the resulting accuracy meaningless.",
        ]),

        MSection("Whole-slide processing and patch embedding", 2, [
            "Slide resolution was read per slide rather than assumed. Tissue was "
            "segmented on a downsampled thumbnail by saturation thresholding with "
            "morphological closing, a simplified surrogate for the published "
            "segmentation step [ref: clam]. Non-overlapping 224x224 patches [PAPER] "
            "were extracted on a regular grid at the target resolution; the published "
            "analysis shows that overlapping strides confer no meaningful benefit. "
            "Patches were screened for pen ink, tissue folds and blur by colour and "
            "focus criteria, and stain-normalised by the Macenko method [ref: macenko] "
            "against a fixed reference matrix so that all patches are normalised to a "
            "common target irrespective of source slide.",
            "Patches were encoded in streaming batches and cached to disk as one "
            "columnar shard per slide, allowing the classification head to be retrained "
            "without re-encoding and allowing an interrupted run to resume without "
            "recomputing completed slides. Each encoder is applied with its own "
            "published preprocessing and normalisation statistics; a generic ImageNet "
            "transform is never substituted, as this degrades embeddings silently "
            "rather than raising an error.",
            "**Departure.** Phikon [ref: phikon] was used as a frozen encoder in place "
            "of the fine-tuned MUSK backbone [ref: musk] of the original study, for "
            "compute reasons and to preserve a frozen-versus-frozen comparison between "
            "Method 1 and Method 2. Slides were processed at 0.50 um per pixel rather "
            "than 0.25 um per pixel; the published analysis reports comparable accuracy "
            "between these resolutions.",
        ]),

        MSection("Habitat classification and the CANVAS-CTX extension", 2, [
            "The classification head follows the published architecture exactly "
            "[PAPER]: a 256-unit layer and a 128-unit layer, each with rectified linear "
            "activation and dropout, followed by a linear map to K+1 logits comprising "
            "the habitat classes and background. Class imbalance is addressed by both "
            "weighted random sampling with weights inversely proportional to class "
            "frequency and focal loss, as specified, rather than by either alone "
            "[PAPER].",
            "CANVAS-CTX augments the input to this unchanged head with a "
            "representation of the index patch's spatial neighbourhood, computed by one "
            "of three encoders. The graph encoder pools the k nearest neighbouring "
            "patches by a deep set with distance-biased attention and is permutation "
            "invariant. The two-dimensional grid encoder arranges neighbours on their "
            "true lattice positions and applies two-dimensional convolution with "
            "learned positional encoding, and is therefore orientation aware. The "
            "multi-scale grid encoder extends this to a scale axis spanning "
            "magnifications, allowing a kernel to combine evidence across resolutions "
            "at nearby positions.",
            "Neighbour indices are constructed strictly within a slide, so context "
            "never crosses slide boundaries. Positions falling outside the slide are "
            "masked, and lattice gaps left by artefact-filtered patches are passed to "
            "the network as an explicit occupancy channel so that absent tissue is "
            "distinguishable from low-signal tissue. Setting the neighbour count to "
            "zero reduces the model exactly to the published per-patch classifier; this "
            "equivalence is enforced by an automated test rather than asserted.",
        ]),

        MSection("Spatial feature engineering", 2, [
            "Each slide's habitat map yields 262 features [PAPER] in six blocks: "
            "habitat composition (10), ecological diversity treating habitats as "
            "species (6), intra-habitat spatial dispersion as planar point patterns (90, "
            "being nine metrics for each of ten habitats), pairwise habitat interaction "
            "scored against a spatial null (100), nearest-neighbour distances within "
            "and between habitats (55), and spatial transition entropy computed over a "
            "patch-level six-nearest-neighbour graph (1) [PAPER]. The total is asserted "
            "at runtime so that a block producing the wrong count fails loudly rather "
            "than silently altering the feature space.",
            "**Departure, interaction null.** Habitat-pair interaction is scored "
            "against a toroidal-shift null rather than by shuffling habitat labels. "
            "Label shuffling destroys all spatial autocorrelation, so the resulting "
            "null describes randomly scattered habitats, a configuration real tissue "
            "never adopts. Tested against it, nearly every habitat pair is declared "
            "significant, and the test effectively asks whether the tissue is "
            "spatially organised at all rather than whether two specific habitats are "
            "associated. The toroidal shift preserves each habitat's own "
            "autocorrelation and domain structure and randomises only the relative "
            "registration between the habitat field and the point pattern. The "
            "magnitude of this difference is quantified in the supplementary material.",
            "**Departure, edge correction.** Ripley's K and L are computed with border "
            "correction and the Clark-Evans index with Donnelly's perimeter correction. "
            "Uncorrected estimators are biased on small tissue regions because "
            "boundary points have artificially few observed neighbours. Where too few "
            "eligible centres remain at a given radius the estimate is returned as "
            "missing rather than as a misleading finite value.",
        ]),

        MSection("Clinical modelling", 2, [
            "Slides were partitioned into tumour bulk and leading edge compartments by "
            "unsupervised spatial clustering of smoothed patch-level tumour "
            "probability and its local gradient, and habitat profiles computed within "
            "each compartment [PAPER]. Univariate Cox proportional hazards models were "
            "fitted per habitat and compartment, reporting hazard ratios per standard "
            "deviation with 95 percent confidence intervals, Wald p values and "
            "Benjamini-Hochberg adjusted q values. Endpoints follow the curated "
            "pan-cancer clinical resource [ref: tcgacdr]. Ecotypes were derived by "
            "consensus clustering of z-scored habitat-by-compartment profiles using "
            "partitioning around medoids with Canberra distance [PAPER], and "
            "characterised by multivariable Cox models adjusted for age, stage and "
            "cancer-appropriate covariates.",
            "The prognostic signature follows the published sequence [PAPER]: "
            "collinear features are reduced by community detection on the Spearman "
            "correlation graph, retaining one representative per community; surviving "
            "features are ranked by selection frequency across resampled LASSO-Cox "
            "fits and independently by random survival forest permutation importance; "
            "and the intersection of the two rankings enters a multivariable Cox model "
            "reporting the concordance index and time-dependent area under the curve.",
            "**Validation caution.** Applying this chain to independently randomised "
            "survival outcomes, univariate testing behaved correctly, returning "
            "essentially no features below the adjusted significance threshold. The "
            "terminal signature model did not: fitted concordance indices of 0.60 to "
            "0.80 were obtained on pure noise, decreasing with increasing sample size. "
            "This is in-sample optimism arising because feature selection and model "
            "evaluation share the same subjects. Concordance indices reported from "
            "this pipeline must therefore derive from held-out data or carry an "
            "explicit optimism correction.",
        ]),

        MSection("Statistics", 2, [
            "All tests are two-sided with alpha of 0.05 and Benjamini-Hochberg control "
            "of the false discovery rate [PAPER]. Two-group comparisons use the t test "
            "or Mann-Whitney U test and multi-group comparisons analysis of variance or "
            "the Kruskal-Wallis test, according to distribution. Classifier performance "
            "is reported as macro-F1 and Cohen's kappa with bootstrap confidence "
            "intervals and per-class recall; overall accuracy is never reported alone, "
            "because the background class dominates and accuracy conceals collapse of a "
            "minority habitat. Model comparisons across context modes use at least six "
            "random seeds, each reshuffling both the slide-level split and model "
            "initialisation, summarised as mean and standard deviation and tested by "
            "paired Wilcoxon signed-rank test. Effect sizes are reported with "
            "dispersion or confidence intervals in all cases; p values are never "
            "reported alone.",
        ]),

        MSection("Data availability", 1, [
            "All primary data are public. Colorectal CODEX single-cell data are "
            "available from Mendeley Data (mpjzbtfgfr) [ref: schurch]. Paired "
            "immunofluorescence and H&E whole-slide images are available from the "
            "public Orion release [ref: orion]. Diagnostic whole-slide images and "
            "clinical outcomes are available from the Genomic Data Commons and the "
            "curated pan-cancer clinical resource [ref: tcgacdr]. Breast imaging mass "
            "cytometry resources referenced for comparison are available from their "
            "respective repositories [ref: danenberg; ref: jackson; ref: keren].",
        ]),

        MSection("Code availability", 1, [
            "Analysis code, configuration files and the automated test suite are "
            "available at [REPOSITORY URL PENDING]. All stages are configuration "
            "driven, seeded and resumable, and each records a manifest of "
            "configuration hash, package versions and input checksums. "
            "*Note for the authors: the working directory is not currently under "
            "version control, so no commit identifier can be cited. Initialise a "
            "repository before submission.*",
        ]),
    ]


# Manuscript figures are composed from the generated report panels. A
# manuscript figure is a plate: several report panels published as one
# numbered display item. Panels that do not exist yet are simply absent from
# the plate, and the legend says which and why.
FIGURE_PANELS: dict[str, list[str]] = {
    "Figure 1": ["F1_study_design", "F2_k_sweep", "F3_cn_marker_heatmap",
                 "F4_cn_spatial_maps", "F5_cn_composition"],
    "Figure 2": ["F6_registration_qc", "F7_patch_labels"],
    "Figure 3": ["F8_training_curves", "F9_confusion_matrix", "F10_benchmark",
                 "F11_params_vs_f1"],
    "Figure 4": ["F13_wsi_habitat_maps", "F12_attention_maps"],
    "Figure 5": ["F14_feature_matrix", "F15_feature_correlation", "F16_cox_forest",
                 "F17_km_tertile", "F18_consensus", "F20_signature"],
    "Supplementary Figure S1": ["F21_null_comparison"],
    "Supplementary Figure S2": ["F22_edge_correction"],
}


def _panel_paths(fig_id: str, figures_dir: str = "figures") -> list[Path]:
    """Existing PNG panels for a manuscript figure, in plate order."""
    return [p for p in (Path(figures_dir) / f"{n}.png"
                        for n in FIGURE_PANELS.get(fig_id, [])) if p.exists()]


def build_figure_legends() -> list[tuple[str, str, str]]:
    """(figure id, title, legend). Legends are written now; the data they will
    describe is not yet acquired, and each says so explicitly."""
    return [
        ("Figure 1", "Study design and cellular neighbourhood discovery",
        "(A) Schematic of the six-stage pipeline, the two model variants and the "
        "substitutions made relative to the published protocol. (B) Selection "
        "diagnostics across candidate cluster numbers, showing silhouette width, "
        "Davies-Bouldin index and adjacent-k adjusted Rand index on shared axes with "
        "the selected solution marked. (C) Marker enrichment per cellular "
        "neighbourhood, z-scored across neighbourhoods. (D) Representative tissue "
        "regions coloured by neighbourhood assignment. (E) Neighbourhood composition "
        "per sample, ordered by dominant neighbourhood. Panels B to E require the "
        "colorectal CODEX single-cell table."),
        ("Figure 2", "Habitat label transfer onto matched H&E",
        "(A) Immunofluorescence and H&E overlay for a representative specimen after "
        "affine alignment. (B) Distribution of centroid-to-centroid registration "
        "residuals in microns with the 5 um acceptance threshold indicated. (C) Patch "
        "label distribution per class and per split, with representative patches for "
        "each habitat. All panels require the paired Orion specimens."),
        ("Figure 3", "Spatial context improves habitat inference",
        "(A) Training loss and validation macro-F1 by epoch for all four context "
        "modes. (B) Row-normalised confusion matrix with per-class recall annotated. "
        "(C) Macro-F1 and Cohen's kappa by context mode, mean and standard deviation "
        "across seeds with individual seeds overlaid and paired test results "
        "annotated; the per-patch baseline is marked as the published reference "
        "method. (D) Parameter count against macro-F1, testing whether any advantage "
        "of the grid encoders reflects capacity rather than spatial structure. All "
        "panels require habitat-labelled embeddings from the Orion cohort."),
        ("Figure 4", "Habitat inference across colorectal and breast cohorts",
        "(A) Whole-slide habitat maps overlaid on H&E with the tumour bulk and "
        "leading edge boundary drawn. (B) Attention weights showing which neighbouring "
        "patches the context model consulted, an output the per-patch model cannot "
        "produce. (C) Habitat composition shift between the training and transfer "
        "cancer types. All panels require diagnostic whole-slide images and a trained "
        "classifier."),
        ("Figure 5", "Spatial features and clinical association",
        "(A) The 262-feature matrix, clustered, annotated by feature block. "
        "(B) Univariate Cox hazard ratios with 95 percent confidence intervals per "
        "habitat and compartment, with false-discovery-significant associations "
        "highlighted. (C) Kaplan-Meier curves by habitat tertile with log-rank tests. "
        "(D) Consensus clustering diagnostics and the resulting ecotypes with clinical "
        "annotation tracks. All panels require habitat maps joined to clinical "
        "outcomes."),
        ("Supplementary Figure S1", "Null-model comparison",
        "Habitat pairs declared significant under a toroidal-shift null against label "
        "shuffling on identical input, with the corresponding distributions of "
        "interaction z-scores. Demonstrates why the toroidal null is used throughout. "
        "Computable without further data acquisition."),
        ("Supplementary Figure S2", "Edge correction in spatial dispersion",
        "Border-corrected and uncorrected Ripley's K across the radius grid against "
        "the expectation under complete spatial randomness, showing the magnitude and "
        "direction of boundary bias. Computable without further data acquisition."),
    ]


def build_table_legends() -> list[tuple[str, str]]:
    return [
        ("Table 1", "Habitat classification performance by context mode: macro-F1 and "
        "Cohen's kappa as mean and standard deviation across seeds, difference from "
        "the per-patch baseline, paired test result, parameter count and training time."),
        ("Table 2", "Univariate Cox associations between habitat abundance and "
        "survival, per compartment: hazard ratio per standard deviation, 95 percent "
        "confidence interval, p value, adjusted q value, sample count and EVENT count."),
        ("Supplementary Table S1", "Dataset inventory: cohort, platform, sample size, "
        "role in the study, accession, and the published resource each replaces."),
        ("Supplementary Table S2", "Cellular neighbourhood definitions: label, assigned "
        "name, most enriched cell types, and frequency."),
        ("Supplementary Table S3", "Patch counts by habitat label and split, with "
        "patient-level split composition."),
        ("Supplementary Table S4", "Classifier performance with bootstrap confidence "
        "intervals."),
        ("Supplementary Table S6", "Per-class precision, recall and F1 for every "
        "habitat and context mode."),
        ("Supplementary Table S7", "Full univariate Cox results across all features."),
        ("Supplementary Table S8", "Ecotype clinical associations with tests used and "
        "adjusted q values."),
        ("Supplementary Table S9", "Prognostic signature: selected features, "
        "coefficients, selection frequency and permutation importance."),
        ("Supplementary Table S10", "Complete itemised list of departures from the "
        "published method: parameter, published value, value used here, rationale and "
        "expected impact."),
    ]


# ==================================================================== rendering


def render_markdown() -> str:
    secs = build_manuscript_sections()
    lines = []
    for s in secs:
        lines.append(f"{'#' * (s.level + 1)} {s.heading}\n")
        for p in s.paragraphs:
            lines.append(p + "\n")

    lines.append("## Figures\n")
    for fid, title, legend in build_figure_legends():
        lines.append(f"### {fid}. {title}\n")
        panels = _panel_paths(fid)
        for p in panels:
            lines.append(f"![{fid} panel]({p.as_posix()})\n")
        if not panels:
            lines.append(f"*[No panel generated yet. {legend.split('.')[-2].strip()}.]*\n")
        lines.append(f"**{fid}. {title}.** {legend}\n")

    lines.append("## Table legends\n")
    for tid, legend in build_table_legends():
        lines.append(f"**{tid}.** {legend}\n")

    lines.append("## References\n")
    for i, r in enumerate(REFERENCES, 1):
        bits = [f"{i}. [{r.key}] {r.citation}"]
        if r.pmid:
            bits.append(f"PMID: {r.pmid}.")
        if r.doi:
            bits.append(f"doi: {r.doi}.")
        if not r.verified:
            bits.append(f"**UNVERIFIED - {r.note}**")
        elif r.note:
            bits.append(f"({r.note})")
        lines.append(" ".join(bits) + "\n")

    lines.append("\n---\n")
    lines.append("*Citation keys appear inline as [ref: key] and must be replaced with "
                "the journal's numbering during submission preparation. Every reference "
                "above was checked against PubMed or the publisher page except where "
                "marked UNVERIFIED.*\n")
    return "\n".join(lines)


def render_docx(out_path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    navy = RGBColor(0x1C, 0x2B, 0x4A)
    steel = RGBColor(0x24, 0x71, 0xA3)
    amber = RGBColor(0xB0, 0x00, 0x20)

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    for s in build_manuscript_sections():
        h = doc.add_heading(s.heading, level=min(s.level, 3))
        for run in h.runs:
            run.font.color.rgb = navy if s.level == 1 else steel
        for p_text in s.paragraphs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(p_text)
            if p_text.startswith("[RESULTS PENDING"):
                run.font.color.rgb = amber
                run.bold = True
            elif p_text.startswith("*"):
                run.italic = True

    h = doc.add_heading("Figures", level=1)
    for run in h.runs:
        run.font.color.rgb = navy
    from docx.shared import Cm
    for fid, title, legend in build_figure_legends():
        sub = doc.add_heading(f"{fid}. {title}", level=3)
        for run in sub.runs:
            run.font.color.rgb = steel
        panels = _panel_paths(fid)
        for panel in panels:
            try:
                doc.add_picture(str(panel), width=Cm(16))
            except Exception:
                doc.add_paragraph(f"[panel could not be embedded: {panel.name}]")
        if not panels:
            warn = doc.add_paragraph()
            r = warn.add_run("[No panel generated yet: this display item awaits the "
                             "data named in its legend.]")
            r.font.color.rgb = amber
            r.bold = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{fid}. {title}. ").bold = True
        p.add_run(legend)

    h = doc.add_heading("Table legends", level=1)
    for run in h.runs:
        run.font.color.rgb = navy
    for tid, legend in build_table_legends():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{tid}. ").bold = True
        p.add_run(legend)

    h = doc.add_heading("References", level=1)
    for run in h.runs:
        run.font.color.rgb = navy
    for i, r in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{i}. ").bold = True
        p.add_run(r.citation + " ")
        if r.pmid:
            p.add_run(f"PMID: {r.pmid}. ")
        if r.doi:
            p.add_run(f"doi: {r.doi}. ")
        if not r.verified:
            warn = p.add_run(f"UNVERIFIED - {r.note}")
            warn.font.color.rgb = amber
            warn.bold = True
        elif r.note:
            p.add_run(f"({r.note})").italic = True

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_manuscript(reports_dir: str = "reports") -> dict:
    out = Path(reports_dir)
    out.mkdir(parents=True, exist_ok=True)
    md = render_markdown()
    (out / "manuscript.md").write_text(md, encoding="utf-8")
    render_docx(str(out / "manuscript.docx"))
    n_pending = md.count("[RESULTS PENDING")
    n_unverified = sum(1 for r in REFERENCES if not r.verified)
    return {"n_pending": n_pending, "n_refs": len(REFERENCES),
           "n_unverified_refs": n_unverified, "reports_dir": str(out)}
